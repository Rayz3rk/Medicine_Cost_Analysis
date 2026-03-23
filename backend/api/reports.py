from fastapi import APIRouter, HTTPException
from fastapi.responses import FileResponse
from docx import Document
from reportlab.pdfgen import canvas
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
import os
import uuid
from backend.db.redis_client import get_redis
import json

router = APIRouter()

REPORT_DIR = "generated_reports"
if not os.path.exists(REPORT_DIR):
    os.makedirs(REPORT_DIR)

def get_session_report(session_id: str):
    redis = get_redis()
    report_raw = redis.hget(f"session:{session_id}", "report")
    if not report_raw:
        raise HTTPException(status_code=404, detail="Report not found or session incomplete")
    return json.loads(report_raw)

@router.get("/reports/{session_id}/pdf")
async def download_pdf(session_id: str):
    report = get_session_report(session_id)
    file_path = os.path.join(REPORT_DIR, f"{session_id}.pdf")
    
    # Simple PDF generation
    c = canvas.Canvas(file_path)
    # Using built-in font for simplicity, Chinese characters might require a custom TTF
    # For robust Chinese support, register a font:
    # pdfmetrics.registerFont(TTFont('SimHei', 'simhei.ttf'))
    # c.setFont('SimHei', 12)
    
    textobject = c.beginText(40, 800)
    textobject.setFont("Helvetica", 12)
    textobject.textLines(f"Cost Analysis Report\nSession: {session_id}\n\n")
    
    textobject.textLines(f"Cost Summary:\n{report.get('cost_summary', '')[:500]}...\n\n")
    textobject.textLines(f"Pricing Strategy:\n{report.get('pricing_strategy', '')[:500]}...\n\n")
    textobject.textLines(f"Supply Chain Advice:\n{report.get('supply_chain_advice', '')[:500]}...")
    
    c.drawText(textobject)
    c.save()
    
    return FileResponse(file_path, filename=f"report_{session_id}.pdf", media_type="application/pdf")

@router.get("/reports/{session_id}/docx")
async def download_docx(session_id: str):
    report = get_session_report(session_id)
    file_path = os.path.join(REPORT_DIR, f"{session_id}.docx")
    
    doc = Document()
    doc.add_heading('Cost Analysis Report', 0)
    
    doc.add_heading('Session ID', level=1)
    doc.add_paragraph(session_id)
    
    doc.add_heading('Cost Summary', level=1)
    doc.add_paragraph(report.get('cost_summary', ''))
    
    doc.add_heading('Pricing Strategy', level=1)
    doc.add_paragraph(report.get('pricing_strategy', ''))
    
    doc.add_heading('Supply Chain Advice', level=1)
    doc.add_paragraph(report.get('supply_chain_advice', ''))
    
    doc.save(file_path)
    
    return FileResponse(file_path, filename=f"report_{session_id}.docx", media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
